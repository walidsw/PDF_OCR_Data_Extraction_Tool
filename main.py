"""
Enhanced PDF OCR Data Extraction Tool with Tkinter
Extracts data from multiple PDFs using OCR and AI, consolidates into reports.
Features: Modern UI, batch export, individual tab exports, bug fixes.

Design:
- UI is fully responsive.
"""

import os
import json
import base64
import threading
import queue
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime

import google.generativeai as genai
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText


class PDFOCRProcessor:
    """Process PDFs and extract structured data using Gemini's vision capabilities."""
    
    def __init__(self, api_key: Optional[str] = None):
        """Initialize the processor with Gemini API client."""
        api_key = api_key or os.environ.get("GOOGLE_API_KEY")
        if not api_key:
            raise ValueError("API key required. Set GOOGLE_API_KEY or provide in GUI.")
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel("gemini-2.5-flash-preview-09-2025")
        self.extracted_data = {}

    def _clean_and_parse_json(self, text: str) -> Dict[str, Any]:
        """Cleans AI response and parses as JSON."""
        json_start = text.find('{')
        json_end = text.rfind('}')
        
        if json_start == -1 or json_end == -1:
            raise json.JSONDecodeError("No JSON found in response.", text, 0)
            
        json_text = text[json_start:json_end+1]
        
        try:
            return json.loads(json_text)
        except json.JSONDecodeError as e:
            raise json.JSONDecodeError(f"Failed to parse JSON: {e.msg}", json_text, e.pos)
    
    def pdf_to_images_base64(self, pdf_path: str, max_pages: int = 10) -> List[str]:
        """Convert PDF pages to base64-encoded images."""
        images_b64 = []
        try:
            pdf_document = fitz.open(pdf_path)
            num_pages = min(len(pdf_document), max_pages)
            
            for page_num in range(num_pages):
                page = pdf_document[page_num]
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                image_data = pix.tobytes("png")
                images_b64.append(base64.standard_b64encode(image_data).decode("utf-8"))
            
            pdf_document.close()
        except Exception as e:
            raise Exception(f"Error processing PDF {pdf_path}: {e}")
        
        return images_b64
    
    def extract_field_notes(self, pdf_path: str) -> Dict[str, Any]:
        """Extract sample numbers and field observations from handwritten field notes."""
        images_b64 = self.pdf_to_images_base64(pdf_path)
        if not images_b64:
            return {}
        
        prompt = """Analyze these handwritten field notes and extract the following information.
                
Please return ONLY valid JSON (no markdown, no code blocks) with this structure:
{
    "samples": [
        {
            "sample_id": "unique sample identifier",
            "sample_type": "type of sample",
            "collection_date": "date collected",
            "collector_name": "person who collected",
            "notes": "any observations or notes"
        }
    ],
    "general_notes": "any overall field notes"
}

If you cannot read certain fields, use null. If there are multiple samples, create multiple entries."""
        
        content = [prompt]
        for img_b64 in images_b64:
            content.append({"data": img_b64, "mime_type": "image/png"})
        
        response = self.model.generate_content(content)
        
        try:
            result = self._clean_and_parse_json(response.text)
            self.extracted_data["field_notes"] = result
            return result
        except json.JSONDecodeError as e:
            raise Exception(f"Error parsing field notes JSON: {e}\nRaw response:\n{response.text}")
    
    def extract_lab_results(self, pdf_path: str) -> Dict[str, Any]:
        """Extract lab results data from PDF."""
        images_b64 = self.pdf_to_images_base64(pdf_path)
        if not images_b64:
            return {}
        
        prompt = """Analyze this lab results document and extract all test data.

Please return ONLY valid JSON (no markdown, no code blocks) with this structure:
{
    "results": [
        {
            "sample_id": "sample identifier",
            "test_name": "name of test performed",
            "result_value": "numeric or text result",
            "unit": "unit of measurement",
            "analysis_date": "date of analysis",
            "status": "status (e.g., Pass/Fail/Complete)"
        }
    ],
    "lab_name": "name of laboratory",
    "report_date": "report generation date"
}

If there are multiple tests per sample, create multiple result entries."""
        
        content = [prompt]
        for img_b64 in images_b64:
            content.append({"data": img_b64, "mime_type": "image/png"})
        
        response = self.model.generate_content(content)
        
        try:
            result = self._clean_and_parse_json(response.text)
            self.extracted_data["lab_results"] = result
            return result
        except json.JSONDecodeError as e:
            raise Exception(f"Error parsing lab results JSON: {e}\nRaw response:\n{response.text}")
    
    def extract_location_notes(self, pdf_path: str) -> Dict[str, Any]:
        """Extract location data from handwritten location notes."""
        images_b64 = self.pdf_to_images_base64(pdf_path)
        if not images_b64:
            return {}
        
        prompt = """Analyze these handwritten location notes and extract location information.

Please return ONLY valid JSON (no markdown, no code blocks) with this structure:
{
    "locations": [
        {
            "sample_id": "sample identifier",
            "latitude": "latitude coordinate or description",
            "longitude": "longitude coordinate or description",
            "location_name": "name/description of location",
            "site_description": "description of sampling site",
            "access_notes": "any access or terrain notes"
        }
    ]
}

If coordinates are not precise, use best estimates or null."""
        
        content = [prompt]
        for img_b64 in images_b64:
            content.append({"data": img_b64, "mime_type": "image/png"})
        
        response = self.model.generate_content(content)
        
        try:
            result = self._clean_and_parse_json(response.text)
            self.extracted_data["location_notes"] = result
            return result
        except json.JSONDecodeError as e:
            raise Exception(f"Error parsing location notes JSON: {e}\nRaw response:\n{response.text}")
    
    def consolidate_data(self) -> pd.DataFrame:
        """Consolidate extracted data from all sources into a unified DataFrame."""
        field_notes = self.extracted_data.get("field_notes", {}).get("samples", [])
        lab_results = self.extracted_data.get("lab_results", {}).get("results", [])
        location_notes = self.extracted_data.get("location_notes", {}).get("locations", [])
        
        consolidated = {}
        
        for note in field_notes:
            sid = note.get("sample_id")
            if sid:
                consolidated[sid] = {
                    "Sample ID": sid,
                    "Sample Type": note.get("sample_type"),
                    "Collection Date": note.get("collection_date"),
                    "Collector": note.get("collector_name"),
                    "Field Notes": note.get("notes"),
                }
        
        for loc in location_notes:
            sid = loc.get("sample_id")
            if sid:
                if sid not in consolidated:
                    consolidated[sid] = {"Sample ID": sid}
                
                consolidated[sid].update({
                    "Location": loc.get("location_name"),
                    "Latitude": loc.get("latitude"),
                    "Longitude": loc.get("longitude"),
                    "Site Description": loc.get("site_description"),
                })
        
        for result in lab_results:
            sid = result.get("sample_id")
            if sid:
                if sid not in consolidated:
                    consolidated[sid] = {"Sample ID": sid}
                
                test_key = f"{result.get('test_name', 'Test')} ({result.get('unit', 'N/A')})"
                consolidated[sid][test_key] = result.get("result_value")
        
        if not consolidated:
            return pd.DataFrame()
            
        df = pd.DataFrame(list(consolidated.values()))
        
        if "Sample ID" in df.columns:
            cols = ["Sample ID"] + [col for col in df.columns if col != "Sample ID"]
            df = df[cols]
            
        return df
    
    def export_to_excel(self, df: pd.DataFrame, output_path: str, sheet_name: str = "Data"):
        """Export DataFrame to Excel."""
        if df is None or df.empty:
            raise ValueError("No data to export")
        df.to_excel(output_path, sheet_name=sheet_name, index=False)
    
    def export_to_word(self, df: pd.DataFrame, output_path: str, title: str = "Data Report"):
        """Export DataFrame to Word document with formatted table."""
        if df is None or df.empty:
            raise ValueError("No data to export")
            
        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        if "lab_results" in self.extracted_data:
            lab_name = self.extracted_data["lab_results"].get("lab_name")
            if lab_name:
                doc.add_paragraph(f"Laboratory: {lab_name}")
        
        doc.add_paragraph()
        doc.add_heading("Results Table", level=1)

        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value) if pd.notna(value) else ""
        
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)

        doc.save(output_path)


class ScrollableFrame(ttk.Frame):
    """A scrollable frame that adjusts to its content."""
    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)
        
        # Get background color from style
        style = ttk.Style()
        bg_color = style.lookup("TFrame", "background")

        self.canvas = tk.Canvas(self, borderwidth=0, highlightthickness=0, background=bg_color)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        
        # This frame will hold the actual content
        self.scrollable_frame = ttk.Frame(self.canvas, style="TFrame")

        # Update scrollregion when the size of scrollable_frame changes
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(
                scrollregion=self.canvas.bbox("all")
            )
        )

        # Put the scrollable_frame inside the canvas
        self.canvas_window = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # Update the width of the scrollable_frame to match the canvas
        self.canvas.bind("<Configure>", self._on_canvas_configure)

        # --- Layout ---
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        
        # Bind mouse wheel scrolling
        self.bind_all("<MouseWheel>", self._on_mousewheel)

    def _on_canvas_configure(self, event):
        # Update the scrollable_frame's width to match the canvas's width
        self.canvas.itemconfig(self.canvas_window, width=event.width)

    def _on_mousewheel(self, event):
        """Handle mouse wheel scrolling."""
        # Determine scroll direction (platform-dependent)
        if event.num == 5 or event.delta < 0:
            scroll_val = 1
        elif event.num == 4 or event.delta > 0:
            scroll_val = -1
        else:
            # Default for Windows-like scrolling (event.delta is 120 or -120)
            scroll_val = -1 * (event.delta // 120)
            
        self.canvas.yview_scroll(scroll_val, "units")


class App(tk.Tk):
    """Main application GUI with modern styling."""
    
    def __init__(self):
        super().__init__()
        
        self.processor: Optional[PDFOCRProcessor] = None
        self.extracted_json: str = ""
        
        self.current_dataframe: Optional[pd.DataFrame] = None
        self.df_field_notes: Optional[pd.DataFrame] = None
        self.df_lab_results: Optional[pd.DataFrame] = None
        self.df_loc_notes: Optional[pd.DataFrame] = None
        
        self.title("PDF OCR Data Extraction Tool")
        self.geometry("1000x750")
        self.minsize(600, 500)  # Set minimum window size
        
        self._setup_styles()
        self.configure(bg=self.bg_color)
        self._create_widgets()
        
        self.queue = queue.Queue()
        self.check_queue()

    def _setup_styles(self):
        """Configure modern ttk styles."""
        style = ttk.Style(self)
        style.theme_use('clam')
        
        # Colors
        self.bg_color = "#f0f0f0"
        self.primary_color = "#0078d4"
        self.accent_color = "#28a745"
        self.text_color = "#333333"
        
        # Configure styles
        style.configure("TFrame", background=self.bg_color)
        style.configure("TLabel", background=self.bg_color, foreground=self.text_color, font=("Segoe UI", 10))
        style.configure("TButton", font=("Segoe UI", 10), padding=6)
        style.map("TButton", background=[("active", self.primary_color), ("disabled", self.bg_color)])
        
        style.configure("Header.TLabel", font=("Segoe UI", 18, "bold"), foreground=self.primary_color)
        style.configure("SubHeader.TLabel", font=("Segoe UI", 12, "bold"), foreground=self.text_color)
        
        style.configure("Primary.TButton", font=("Segoe UI", 10, "bold"), padding=10)
        style.map("Primary.TButton", 
                  background=[("active", "#005a9e"), ("!disabled", self.primary_color)],
                  foreground=[("!disabled", "white")])

        style.configure("Success.TButton", font=("Segoe UI", 10), padding=8)
        style.map("Success.TButton", 
                  background=[("active", "#1e7e34"), ("!disabled", self.accent_color)],
                  foreground=[("!disabled", "white")])
        
        style.configure("TNotebook", background=self.bg_color, borderwidth=0)
        style.configure("TNotebook.Tab", font=("Segoe UI", 9), padding=[10, 5])
        style.map("TNotebook.Tab", 
                  background=[("selected", self.bg_color)],
                  expand=[("selected", [1, 1, 1, 0])])
        
        style.configure("TLabelframe", background=self.bg_color, borderwidth=1, relief="solid")
        style.configure("TLabelframe.Label", background=self.bg_color, foreground=self.text_color, font=("Segoe UI", 10, "bold"))

    def _create_widgets(self):
        """Create all GUI widgets with responsive layout."""
        
        # --- 1. Header (Fixed at Top) ---
        header_frame = ttk.Frame(self, padding=(15, 15, 15, 10))
        header_frame.pack(fill="x", side="top")
        
        ttk.Label(header_frame, text="📄 PDF OCR Data Extraction Tool", 
                  style="Header.TLabel").pack(side="left")
        
        ttk.Separator(self, orient="horizontal").pack(fill="x", side="top", padx=15)
        
        # --- 3. Action Bar (Fixed at Bottom) ---
        self._create_action_bar()
        
        # --- 2. Scrollable Content (Fills Middle) ---
        self.scrollable_area = ScrollableFrame(self)
        self.scrollable_area.pack(fill="both", expand=True, side="top", padx=15, pady=(5, 0))
        
        # content_frame is the inner frame to pack widgets into
        content_frame = self.scrollable_area.scrollable_frame
        content_frame.columnconfigure(0, weight=1)

        # Create sections inside the scrollable content_frame
        self._create_file_section(content_frame)
        self._create_api_section(content_frame)
        self._create_results_section(content_frame)

    def _create_file_section(self, parent):
        """Create file selection section."""
        file_frame = ttk.LabelFrame(parent, text="  1️⃣ Select PDF Files  ", padding="10")
        file_frame.pack(fill="x", expand=True, pady=5)
        file_frame.columnconfigure(1, weight=1)
        
        self.field_notes_var = tk.StringVar()
        self.lab_results_var = tk.StringVar()
        self.location_notes_var = tk.StringVar()
        
        self._create_file_row(file_frame, "📝 Field Notes:", self.field_notes_var, 0)
        self._create_file_row(file_frame, "🔬 Lab Results:", self.lab_results_var, 1)
        self._create_file_row(file_frame, "📍 Location Notes:", self.location_notes_var, 2)

    def _create_file_row(self, parent, label_text, string_var, row):
        """Create a file selection row with grid layout."""
        ttk.Label(parent, text=label_text, width=18).grid(row=row, column=0, padx=5, pady=5, sticky="w")
        
        entry = ttk.Entry(parent, textvariable=string_var, state="readonly")
        entry.grid(row=row, column=1, padx=5, pady=5, sticky="ew")
        
        def browse():
            filepath = filedialog.askopenfilename(
                title=f"Select {label_text}",
                filetypes=(("PDF Files", "*.pdf"), ("All Files", "*.*"))
            )
            if filepath:
                string_var.set(filepath)
        
        ttk.Button(parent, text="Browse", command=browse, width=10).grid(
            row=row, column=2, padx=5, pady=5
        )

    def _create_api_section(self, parent):
        """Create API configuration section."""
        api_frame = ttk.LabelFrame(parent, text="  2️⃣ API Configuration  ", padding="10")
        api_frame.pack(fill="x", expand=True, pady=5)
        api_frame.columnconfigure(1, weight=1)
        
        ttk.Label(api_frame, text="🔑 Google Gemini API Key:").grid(
            row=0, column=0, padx=5, pady=5, sticky="w"
        )
        
        self.api_key_entry = ttk.Entry(api_frame, show="●")
        self.api_key_entry.insert(0, os.environ.get("GOOGLE_API_KEY", ""))
        self.api_key_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        
        ttk.Label(api_frame, text="(Leave empty to use GOOGLE_API_KEY environment variable)",
                  font=("Segoe UI", 8), foreground="gray").grid(
            row=1, column=0, columnspan=2, padx=5, pady=(0, 5), sticky="w"
        )

    def _create_results_section(self, parent):
        """Create results display section."""
        results_frame = ttk.LabelFrame(parent, text="  3️⃣ Results & Preview  ", padding="10")
        results_frame.pack(fill="both", expand=True, pady=5)
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(3, weight=1)
        
        self.progress_bar = ttk.Progressbar(results_frame, mode="determinate")
        self.progress_bar.grid(row=0, column=0, sticky="ew", pady=(0, 5))
        
        status_label = ttk.Label(results_frame, text="Status Log:", font=("Segoe UI", 9, "bold"))
        status_label.grid(row=1, column=0, sticky="w")
        
        self.status_text = ScrolledText(results_frame, height=5, state="disabled", 
                                        wrap="word", font=("Consolas", 9))
        self.status_text.grid(row=2, column=0, sticky="ew", pady=(0, 5))
        
        self.notebook = ttk.Notebook(results_frame)
        self.notebook.grid(row=3, column=0, sticky="nsew", pady=5)
        
        self.field_notes_preview_text = self._create_preview_tab("📝 Field Notes")
        self.lab_results_preview_text = self._create_preview_tab("🔬 Lab Results")
        self.location_notes_preview_text = self._create_preview_tab("📍 Location Notes")
        self.consolidated_preview_text = self._create_preview_tab("📊 Consolidated")
        self.full_json_text = self._create_preview_tab("📋 Full JSON", wrap="word")

    def _create_preview_tab(self, title, wrap="none"):
        """Create a preview tab."""
        frame = ttk.Frame(self.notebook, padding=0)
        text_widget = ScrolledText(frame, state="disabled", wrap=wrap, 
                                   font=("Consolas", 9), borderwidth=0, highlightthickness=0)
        text_widget.pack(fill="both", expand=True)
        self.notebook.add(frame, text=title)
        return text_widget

    def _create_action_bar(self):
        """Create bottom action button bar."""
        action_frame = ttk.Frame(self, padding=(15, 10, 15, 15))
        action_frame.pack(fill="x", side="bottom")
        
        ttk.Separator(action_frame, orient="horizontal").pack(fill="x", expand=True, pady=(0, 10))
        
        button_container = ttk.Frame(action_frame)
        button_container.pack()
        
        self.submit_button = ttk.Button(
            button_container, text="✅ SUBMIT & PROCESS PDFs", 
            command=self.start_processing,
            style="Primary.TButton",
            width=30
        )
        self.submit_button.grid(row=0, column=0, columnspan=4, pady=(0, 10), padx=5)
        
        ttk.Label(button_container, text="Export Options:", 
                 font=("Segoe UI", 9, "bold")).grid(row=1, column=0, padx=5, sticky="e")
        
        self.download_all_button = ttk.Button(
            button_container, text="💾 Download All", 
            command=self.download_all_excel,
            state="disabled",
            style="Success.TButton",
            width=15
        )
        self.download_all_button.grid(row=1, column=1, padx=5)
        
        self.export_excel_button = ttk.Button(
            button_container, text="📊 Export Excel", 
            command=self.export_to_excel,
            state="disabled",
            width=15
        )
        self.export_excel_button.grid(row=1, column=2, padx=5)
        
        self.export_word_button = ttk.Button(
            button_container, text="📄 Export Word", 
            command=self.export_to_word,
            state="disabled",
            width=15
        )
        self.export_word_button.grid(row=1, column=3, padx=5)

    def update_status(self, msg, clear=False):
        """Update status text box."""
        self.status_text.config(state="normal")
        if clear:
            self.status_text.delete("1.0", tk.END)
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.status_text.insert(tk.END, f"[{timestamp}] {msg}\n")
        self.status_text.see(tk.END)
        self.status_text.config(state="disabled")

    def start_processing(self):
        """Validate inputs and start processing thread."""
        field_notes_path = self.field_notes_var.get()
        lab_results_path = self.lab_results_var.get()
        location_notes_path = self.location_notes_var.get()
        api_key = self.api_key_entry.get().strip()

        if not all([field_notes_path, lab_results_path, location_notes_path]):
            messagebox.showerror("Missing Files", "Please select all three PDF files.")
            return
            
        if not (api_key or os.environ.get("GOOGLE_API_KEY")):
            messagebox.showerror("Missing API Key", "Please provide a Google Gemini API Key.")
            return

        # Disable buttons
        self.submit_button.config(state="disabled")
        self.export_excel_button.config(state="disabled")
        self.export_word_button.config(state="disabled")
        self.download_all_button.config(state="disabled")

        # Clear previous results
        for widget in [self.field_notes_preview_text, self.lab_results_preview_text,
                       self.location_notes_preview_text, self.consolidated_preview_text,
                       self.full_json_text]:
            widget.config(state="normal")
            widget.delete("1.0", tk.END)
            widget.config(state="disabled")

        self.current_dataframe = None
        self.df_field_notes = None
        self.df_lab_results = None
        self.df_loc_notes = None
        self.extracted_json = ""
        
        # Start processing thread
        thread = threading.Thread(
            target=self.processing_thread_task,
            args=(api_key, field_notes_path, lab_results_path, location_notes_path),
            daemon=True
        )
        thread.start()

    def processing_thread_task(self, api_key, field_path, lab_path, loc_path):
        """Processing work in separate thread."""
        try:
            self.queue.put(("status", "Initializing processor...", True))
            self.processor = PDFOCRProcessor(api_key=api_key)
            self.queue.put(("progress", 10))

            self.queue.put(("status", "Extracting field notes..."))
            self.processor.extract_field_notes(field_path)
            self.queue.put(("progress", 30))
            
            self.queue.put(("status", "Extracting lab results..."))
            self.processor.extract_lab_results(lab_path)
            self.queue.put(("progress", 60))
            
            self.queue.put(("status", "Extracting location data..."))
            self.processor.extract_location_notes(loc_path)
            self.queue.put(("progress", 80))
            
            self.queue.put(("status", "Consolidating data..."))
            self.current_dataframe = self.processor.consolidate_data()
            self.extracted_json = json.dumps(self.processor.extracted_data, indent=2)
            self.queue.put(("progress", 100))

            self.queue.put(("status", "✅ Processing complete!"))
            self.queue.put(("done", None))

        except Exception as e:
            self.queue.put(("error", str(e)))

    def check_queue(self):
        """Check queue for thread messages."""
        try:
            while True:
                msg_type, msg_data = self.queue.get_nowait()
                
                if msg_type == "status":
                    if isinstance(msg_data, tuple):
                        text, clear = msg_data[0], msg_data[1]
                    else:
                        text, clear = msg_data, False
                    self.update_status(text, clear=clear)
                
                elif msg_type == "progress":
                    self.progress_bar.config(value=msg_data)
                
                elif msg_type == "done":
                    self.handle_processing_done()
                
                elif msg_type == "error":
                    self.handle_processing_error(msg_data)
                    
        except queue.Empty:
            pass
        finally:
            self.after(100, self.check_queue)

    def handle_processing_done(self):
        """Handle completion of processing."""
        self.progress_bar.config(value=0)
        self.submit_button.config(state="normal")
        
        if not self.processor:
            return

        def populate_tab(text_widget, data_list, empty_msg):
            text_widget.config(state="normal")
            text_widget.delete("1.0", tk.END)
            df = None
            if data_list:
                df = pd.DataFrame(data_list)
                text_widget.insert("1.0", df.to_string())
            else:
                text_widget.insert("1.0", empty_msg)
            text_widget.config(state="disabled")
            return df

        field_data = self.processor.extracted_data.get("field_notes", {}).get("samples", [])
        self.df_field_notes = populate_tab(self.field_notes_preview_text, field_data, 
                                          "No field notes extracted.")
        
        lab_data = self.processor.extracted_data.get("lab_results", {}).get("results", [])
        self.df_lab_results = populate_tab(self.lab_results_preview_text, lab_data, 
                                          "No lab results extracted.")
        
        loc_data = self.processor.extracted_data.get("location_notes", {}).get("locations", [])
        self.df_loc_notes = populate_tab(self.location_notes_preview_text, loc_data, 
                                        "No location notes extracted.")

        if self.extracted_json:
            self.full_json_text.config(state="normal")
            self.full_json_text.delete("1.0", tk.END)
            self.full_json_text.insert("1.0", self.extracted_json)
            self.full_json_text.config(state="disabled")
            
        self.consolidated_preview_text.config(state="normal")
        self.consolidated_preview_text.delete("1.0", tk.END)
        if self.current_dataframe is not None and not self.current_dataframe.empty:
            self.consolidated_preview_text.insert("1.0", self.current_dataframe.to_string())
            self.export_excel_button.config(state="normal")
            self.export_word_button.config(state="normal")
            self.download_all_button.config(state="normal")
        else:
            self.update_status("⚠️ Warning: No data was consolidated.")
            self.consolidated_preview_text.insert("1.0", "No consolidated data available.")
        self.consolidated_preview_text.config(state="disabled")

    def handle_processing_error(self, error_msg):
        """Handle processing errors."""
        self.progress_bar.config(value=0)
        self.submit_button.config(state="normal")
        self.update_status(f"❌ Error: {error_msg}")
        messagebox.showerror("Processing Error", f"An error occurred:\n\n{error_msg}")

    def _get_active_dataframe(self):
        """Get DataFrame for currently active tab."""
        try:
            selected_index = self.notebook.index(self.notebook.select())
        except tk.TclError:
            return None, None

        tab_map = {
            0: (self.df_field_notes, "Field Notes"),
            1: (self.df_lab_results, "Lab Results"),
            2: (self.df_loc_notes, "Location Notes"),
            3: (self.current_dataframe, "Consolidated Report"),
            4: (None, "JSON")
        }
        
        return tab_map.get(selected_index, (None, None))

    def download_all_excel(self):
        """Download all 3 individual files + consolidated in separate Excel files."""
        if not self.processor:
            messagebox.showwarning("No Data", "Please process PDFs first.")
            return
        
        folder = filedialog.askdirectory(title="Select folder to save all Excel files")
        if not folder:
            return
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_files = []
        
        try:
            # Export each individual dataset
            exports = [
                (self.df_field_notes, f"field_notes_{timestamp}.xlsx", "Field Notes"),
                (self.df_lab_results, f"lab_results_{timestamp}.xlsx", "Lab Results"),
                (self.df_loc_notes, f"location_notes_{timestamp}.xlsx", "Location Notes"),
                (self.current_dataframe, f"consolidated_report_{timestamp}.xlsx", "Consolidated")
            ]
            
            for df, filename, sheet_name in exports:
                if df is not None and not df.empty:
                    filepath = os.path.join(folder, filename)
                    self.processor.export_to_excel(df, filepath, sheet_name)
                    saved_files.append(filename)
                    self.update_status(f"✅ Saved: {filename}")
            
            if saved_files:
                file_list = "\n".join([f"  • {f}" for f in saved_files])
                messagebox.showinfo(
                    "Download Complete", 
                    f"Successfully saved {len(saved_files)} file(s):\n\n{file_list}\n\nLocation: {folder}"
                )
            else:
                messagebox.showwarning("No Data", "No data available to export.")
                
        except Exception as e:
            messagebox.showerror("Export Failed", f"Failed to export files:\n\n{str(e)}")

    def export_to_excel(self):
        """Export current tab's data to Excel."""
        df, tab_name = self._get_active_dataframe()
        
        if df is None or df.empty:
            messagebox.showwarning("No Data", "No data available on current tab to export.")
            return
        
        default_name = f"{tab_name.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            initialfile=default_name,
            filetypes=(("Excel Files", "*.xlsx"), ("All Files", "*.*"))
        )
        
        if filepath and self.processor:
            try:
                self.processor.export_to_excel(df, filepath, tab_name)
                self.update_status(f"✅ Exported to Excel: {os.path.basename(filepath)}")
                messagebox.showinfo("Success", f"Excel file saved:\n{filepath}")
            except Exception as e:
                messagebox.showerror("Export Failed", f"Export failed:\n\n{str(e)}")
        
    def export_to_word(self):
        """Export current tab's data to Word."""
        df, tab_name = self._get_active_dataframe()
        
        if df is None or df.empty:
            messagebox.showwarning("No Data", "No data available on current tab to export.")
            return
        
        default_name = f"{tab_name.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            initialfile=default_name,
            filetypes=(("Word Documents", "*.docx"), ("All Files", "*.*"))
        )
        
        if filepath and self.processor:
            try:
                self.processor.export_to_word(df, filepath, tab_name)
                self.update_status(f"✅ Exported to Word: {os.path.basename(filepath)}")
                messagebox.showinfo("Success", f"Word document saved:\n{filepath}")
            except Exception as e:
                messagebox.showerror("Export Failed", f"Export failed:\n\n{str(e)}")


if __name__ == "__main__":
    app = App()
    app.mainloop()

